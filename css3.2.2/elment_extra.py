import re
from docx import Document
from pandas.core.frame import DataFrame

def re_extract(filename):
    doc = Document(filename)
    elements = []
    length=len(doc.paragraphs)
    cache=doc.paragraphs
#0
    # 匹配合同编号
    aa = 0
    elements_index = ['合同编号']
    for i in range(length):
        if re.findall(r'.*(?:合同编号|编号)[：:](.*)号',cache[i].text):
            a = re.search(r'.*(?:合同编号|编号)[：:](.*)号', cache[i].text)
            elements.append(a.group(1))
            aa = 1
            break
    if aa == 0:
        elements.append('')

# 1
    # 适用于
    bb = 0
    elements_index.append('适用于')
    for i in range(length):
        if re.match(r'(.*)协议', cache[i].text):
            a = re.search(r'(.*)协议', cache[i].text)
            elements.append(a.group(1))
            bb = 1
            break
    if bb == 0:
        elements.append('')
# 2
    # 协议签订日期
    c = 0
    elements_index.append('协议签订日期')
    for i in range(length):
        if re.match(r'.*协议签订日期[：:](.*)', cache[i].text):
            a = re.search(r'.*协议签订日期[：:](.*)', cache[i].text)
            elements.append(a.group(1))
            c = 1
            break
    if c == 0:
        elements.append('')
# 3
    # 协议签订地点
    d = 0
    elements_index.append('协议签订地点')
    for i in range(length):
        if re.match(r'.*协议签订地点[：:](.*)', cache[i].text):
            a = re.search(r'.*协议签订地点[：:](.*)', cache[i].text)
            elements.append(a.group(1))
            d = 1
            break
    if d == 0:
        elements.append('')
# 4
    # 转让方
    pattern = '.*转 *让 *方.*[：:](.*)'
    e = 0
    elements_index.append('转让方')
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            elements.append(a.group(1))
            e = 1
            break
    if e == 0:
        elements.append('')
# 5
    # 受让方
    f = 0
    elements_index.append('受让方')
    for i in range(length):
        if re.match(r'.*受 *让 *方.*[：:](.*)', cache[i].text):
            a = re.search(r'.*受 *让 *方.*[：:](.*)', cache[i].text)
            elements.append(a.group(1))
            f = 1
            break
    if f == 0:
        elements.append('')
#6
    # 债务方
    f = 0
    elements_index.append('债务方')
    for i in range(length):
        if re.match(r'.*(?:债务方|债务人)[：:](.*)', cache[i].text):
            a = re.search(r'.*(?:债务方|债务人)[：:](.*)', cache[i].text)
            elements.append(a.group(1))
            f = 1
            break
    if f == 0:
        elements.append('')
# 7
    # 负责人
    g = 0
    ggg = 0
    elements_index.append('转让方负责人')
    pattern = '.*转 *让 *方.*[：:](.*)'
    pattern_1 = '(.*)[：:](.*)'
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            g = 1
            if re.match(pattern_1,cache[i+1].text):
                bsk=re.search(pattern_1,cache[i+1].text)
                ggg = 1
                elements.append(bsk.group(2))
                ggg = 1
                break
    if g == 0:
        elements.append('')
# 8
    # 受让方负责人
    f = 0
    fff = 0
    elements_index.append('受让方负责人')
    pattern = '.*受 *让 *方.*[：:](.*)'
    pattern_1 = '(.*)[：:](.*)'
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            f = 1
            if re.match(pattern_1,cache[i+1].text):
                bsk=re.search(pattern_1,cache[i+1].text)
                fff = 1
                elements.append(bsk.group(2))
                fff = 1
                break
    if f == 0:
        elements.append('')
#9
    # 债务方负责人
    f = 0
    fff = 0
    elements_index.append('债务方负责人')
    pattern = '.*(?:债务方|债务人)[：:](.*)'
    pattern_1 = '(.*)[：:](.*)'
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            f = 1
            if re.match(pattern_1,cache[i+1].text):
                bsk=re.search(pattern_1,cache[i+1].text)
                fff = 1
                elements.append(bsk.group(2))
                fff = 1
                break
    if f == 0:
        elements.append('')

# 10
    # 住所
    ii = 0
    iii = 0
    elements_index.append('转让方住所')
    pattern = '.*转 *让 *方.*[：:](.*)'
    pattern_1 = '(.*)[：:](.*)'
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            ii = 1
            if re.match(pattern_1, cache[i + 2].text):
                bsk = re.search(pattern_1, cache[i + 2].text)
                iii = 1
                elements.append(bsk.group(2))
                iii = 1
                break
    if ii == 0:
        elements.append('')

# 11
    # 地址
    ii = 0
    iii = 0
    elements_index.append('受让方住所')
    pattern = '.*受 *让 *方.*[：:](.*)'
    pattern_1 = '(.*)[：:](.*)'
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            ii = 1
            if re.match(pattern_1, cache[i + 2].text):
                bsk = re.match(pattern_1, cache[i + 2].text)
                iii = 1
                elements.append(bsk.group(2))
                iii = 1
                break
    if ii == 0:
        elements.append('')
#12
    # 债务方住所
    f = 0
    fff = 0
    elements_index.append('债务方住所')
    pattern = '.*(?:债务方|债务人)[：:](.*)'
    pattern_1 = '(.*)[：:](.*)'
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            f = 1
            if re.match(pattern_1,cache[i+2].text):
                bsk=re.search(pattern_1,cache[i+2].text)
                fff = 1
                elements.append(bsk.group(2))
                fff = 1
                break
    if f == 0:
        elements.append('')
# 13
    # 账面本金余额 利息 其他债权 整体债权
    k = 0
    elements_index.append('账面本金余额')
    for para in doc.paragraphs:
        if re.match(r'.*小写[:：](.*)元.*小写[:：](.*)元.*其他债权.*小写[:：](.*)元.*整体债权.*小写[:：](.*)元', para.text):
            a = re.search(r'.*小写[:：](.*)元.*小写[:：](.*)元.*其他债权.*小写[:：](.*)元.*整体债权.*小写[:：](.*)元', para.text)
            elements.append(a.group(1))
            k = 1
            break
    if k == 0:
        elements.append("")

#14
    #利息
    kkk = 0
    elements_index.append('利息')
    for para in doc.paragraphs:
        if re.search(r'.*小写[:：](.*)元.*小写[:：](.*)元.*其他债权.*小写[:：](.*)元.*整体债权.*小写[:：](.*)元', para.text):
            a = re.search(r'.*小写[:：](.*)元.*小写[:：](.*)元.*其他债权.*小写[:：](.*)元.*整体债权.*小写[:：](.*)元', para.text)
            elements.append(a.group(2))
            kkk = 1
            break
    if kkk == 0:
        elements.append("")

#15
    #其他债权
    zzz = 0
    elements_index.append('其他债权')
    for para in doc.paragraphs:
        if re.match(r'.*其他债权.*小写[:：](.*)元.*，.', para.text):
            a = re.search(r'.*其他债权.*小写[:：](.*)元.*，.', para.text)
            elements.append(a.group(1))
            zzz = 1
            break
    if zzz == 0:
        elements.append("")

# 16
    #整体债权
    zzzz = 0
    elements_index.append('整体债权')
    for para in doc.paragraphs:
        if re.match(r'.*整体债权.*小写[:：](.*)元.*.', para.text):
            a = re.search(r'.*整体债权.*小写[:：](.*)元.*.', para.text)
            elements.append(a.group(1))
            zzzz = 1
            break
    if zzzz == 0:
        elements.append("")

#17
    #转让价款
    gg = 0
    elements_index.append('转让价款')
    for i in range(length):
        if re.search(r'.*转让价款为(.*)', cache[i].text):
            a = re.search(r'.*转让价款为(.*)', cache[i].text)
            elements.append(a.group(1))
            gg = 1
            break
    if gg == 0:
        elements.append("")

# 18
    # 债权金额
    nnn = 0
    elements_index.append('债权金额')
    for i in range(length):
        if re.search(r'.*债权金额为.*(.*)', cache[i].text):
            a = re.search(r'.*债权金额为(.*)', cache[i].text)
            elements.append(a.group(1))
            nnn = 1
            break
    if nnn == 0:
        elements.append('')
# 19
    # 违约金
    nnn = 0
    elements_index.append('违约金')
    for i in range(length):
        if re.search(r'(.*)违约金.*小写[:：](.*)元.*.', cache[i].text):
            a = re.search(r'(.*)违约金.*小写[:：](.*)元.*.', cache[i].text)
            elements.append(a.group(2))
            nnn = 1
            break
    if nnn == 0:
        elements.append('')
# 20
    # 基准日_非收购处置类
    nnn = 0
    elements_index.append('基准日_非标')
    for i in range(length):
        if re.search(r'.*“基准日”系(.*)', cache[i].text):
            a = re.search(r'.*“基准日”系(.*)', cache[i].text)
            elements.append(a.group(1))
            nnn = 1
            break
    if nnn == 0:
        elements.append('')
#21
# 基准日_收购处置类
    nnn = 0
    elements_index.append('基准日_标准')
    for i in range(length):
        if re.search(r'.*基准日.*指(?:甲方|转让方)确定的计算标的债权账面本金及利息余额的截止日，(.*)', cache[i].text):
            a = re.search(r'.*基准日.*指(?:甲方|转让方)确定的计算标的债权账面本金及利息余额的截止日，(.*)', cache[i].text)
            elements.append(a.group(1))
            nnn = 1
            break
    if nnn == 0:
        elements.append('')
#22
    # 本息总额_收购处置合同
    kkk = 0
    elements_index.append('本息总额')
    for para in doc.paragraphs:
        if re.search(r'.*本息总额.*小写[:：](.*)元.*，.*小写[:：](.*)元.*，.*小写[:：](.*)元.*', para.text):
            a = re.search(r'.*本息总额.*小写[:：](.*)元.*，.*小写[:：](.*)元.*，.*小写[:：](.*)元.*', para.text)
            elements.append(a.group(1))
            kkk = 1
            break
    if kkk == 0:
        elements.append("")
#23
    # 本金余额_收购处置合同
    kkk = 0
    elements_index.append('本金余额')
    for para in doc.paragraphs:
        if re.search(r'.*本息总额.*小写[:：](.*)元.*，.*小写[:：](.*)元.*，.*小写[:：](.*)元.*', para.text):
            a = re.search(r'.*本息总额.*小写[:：](.*)元.*，.*小写[:：](.*)元.*，.*小写[:：](.*)元.*', para.text)
            elements.append(a.group(2))
            kkk = 1
            break
    if kkk == 0:
        elements.append("")
#24
    # 本息总额_收购处置合同
    kkk = 0
    elements_index.append('欠息')
    for para in doc.paragraphs:
        if re.search(r'.*本息总额.*小写[:：](.*)元.*，.*本金余额.*小写[:：](.*)元.*小写[:：](.*)元.*', para.text):
            a = re.search(r'.*本息总额.*小写[:：](.*)元.*，.*本金余额.*小写[:：](.*)元.*小写[:：](.*)元.*', para.text)
            elements.append(a.group(3))
            kkk = 1
            break
    if kkk == 0:
        elements.append("")
# 25
    # 开户银行
    ll = 0
    elements_index.append('转让方开户银行')
    for i in range(length):
        if re.match(r'.*(?:开户银行|开户行)[：:](.*)', cache[i].text):
            a = re.search(r'.*(?:开户银行|开户行)[：:](.*)', cache[i].text)
            elements.append(a.group(1))
            ll = 1
            break
    if ll == 0:
        elements.append('')
#26
    # 户名
    m = 0
    elements_index.append('转让方户名')
    for i in range(length):
        if re.match(r'.*户 *名[：:](.*)', cache[i].text):
            a = re.search(r'.*户 *名[：:](.*)', cache[i].text)
            elements.append(a.group(1))
            m = 1
            break
    if m == 0:
        elements.append('')
#27
    # 账户
    ii = 0
    iii = 0
    elements_index.append('转让方账号')
    pattern = '.*户 *名[：:](.*)'
    pattern_1 = '(.*)[：:](.*)'
    for i in range(length):
        if re.match(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            ii = 1
            if re.match(pattern_1, cache[i + 1].text):
                bsk = re.match(pattern_1, cache[i + 1].text)
                iii = 1
                elements.append(bsk.group(2))
                iii = 1
                break
    if ii == 0:
        elements.append('')
# 28
    # 交易保证金
    ii = 0
    iii = 0
    elements_index.append('交易保证金详情')
    pattern = '交易保证金'
    pattern_1 = '(.*)，(.*)'
    for i in range(length):
        if re.search(pattern, cache[i].text):
            a = re.search(pattern, cache[i].text)
            ii = 1
            if re.match(pattern_1, cache[i + 1].text):
                bsk = re.search(pattern_1, cache[i+1].text)
                elements.append(bsk.group(2))
                iii = 1
                break
    if ii == 0:
        elements.append('')

    dic = {'element_index': elements_index, 'elements': elements}

    print(dic)
    elements_d = DataFrame(dic)
    print(elements_d)
    return dic




# if __name__ == '__main__':
#     re_extract('我司对外转让_sgcz_zqzr.docx')
