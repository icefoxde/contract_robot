import os
import re
import random
import xsh_exam
import datetime
import pandas as pd
from flask import session
from docx import Document
from docx.oxml.ns import qn
from models import User,Sugg
from config import DevConfig
from models import docx_table
from docx.shared import Inches
from flask_script import Manager
from flask_bootstrap import Bootstrap
from flask_login import login_required
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from flask import Flask, request, current_app
from flask_login import UserMixin, AnonymousUserMixin, current_user
from contract_generation_v3 import contract_generation,save_para_num_txt
from flask_principal import identity_changed, Identity, AnonymousIdentity
from jzs_list_tools import strings_to_list, list_to_dict, get_elements_list
from flask_login import LoginManager, login_user, logout_user, login_required
from flask import render_template, redirect, url_for, send_from_directory, flash
from element_extract_from_txtanddocx_v3 import d_x_dic, check_res, output_elements_from_docx


app = Flask(__name__)
# Bootstrap 的初始化方法
bootstrap = Bootstrap(app)
# 配置数据库连接
app.config.from_object(DevConfig)
app.config['BOOTSTRAP_SERVE_LOCAL'] = True
# 该配置为True，则每次请求都会自动commit数据库的变动
app.config['SQLALCHEMY_COMMIT_ON_TEARDOWN'] = True
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = True

app.config['SECRET_KEY'] = os.urandom(24)

db = SQLAlchemy(app)
manager = Manager(app)

# 定义用户登陆
login_manager = LoginManager()
login_manager.session_protection = "strong"
login_manager.login_view = "login"
login_manager.login_message = ""
login_manager.login_message_category = "info"


@login_manager.user_loader
def load_user(userId):
    try:
        return User.query.filter_by(user_id=userId).first()
    except:
        db.session.rollback()
        raise
    finally:
        db.session.close()

login_manager.init_app(app)

# ----------------------------------------------------------------------------
# # ------------------------------------------------------------------------
# # -------------------------------------------------------------------------
# 后台首页-->登陆页面
@app.route('/')
def hello_world():
    save_para_num_txt()
    return render_template('login.html')


# 后台登陆-->登陆验证
@app.route('/login', methods=['POST', 'GET'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(user_name=username, user_pass=password).one_or_none()
        if username == '':
            flash('请输入用户名')
            return render_template('login.html')
        elif password == '':
            flash('请输入密码')
            return render_template('login.html')
        elif user is not None:
            login_user(user, remember=True)
            identity_changed.send(current_app._get_current_object(), identity=Identity(user.user_id))
            return redirect(url_for('pagecs'))
        else:
            flash('用户名或密码不正确')
            return render_template('login.html')
        db.session.close()
    return render_template('login.html')


# 合同页面选择
@app.route('/pagecs', methods=['POST', 'GET'])
@login_required
def pagecs():
    return render_template('pagechoose.html')


# 后台首页
@app.route('/admin_index', methods=['POST', 'GET'])
@login_required
def admin_index():
    return render_template('admin_index.html')


# 合同审核 上传合同文档 将文档保存到 static/kuploads 文件夹下
# 保存的文件名是 filename
@app.route('/admin_exam', methods=['POST', 'GET'])
@login_required
def admin_exam():
    if request.method == 'POST':
        f = request.files['file']
        basepath = os.path.dirname(__file__)
        upload_path = os.path.join(basepath, 'static/kuploads', secure_filename(f.filename))
        filename = secure_filename(f.filename)
        f.save(upload_path)
        return redirect(url_for('admin_examdoc', filename=filename))
    return render_template('admin_index.html')


# 合同审核后的内容 形式化审核 审核信息 合同流程表
@app.route('/admin_exam/<string:filename>', methods=['POST', 'GET'])
@login_required
def admin_examdoc(filename):
    try:
        # 读取上传后的docx文档
        basepath = os.path.dirname(__file__)
        upload_path_userup = os.path.join(basepath, 'static/kuploads', filename)
        document = Document(upload_path_userup)
        filename = 'docxexam_userupload.docx'
        basepath = os.path.dirname(__file__)
        upload_path_svup = os.path.join(basepath, 'static/docx_repo/docx_gt/', filename)
        document.save(upload_path_svup)

    except:
        flash("请上传后缀为.docx的文件")
        flash("如“华融合同.docx”")
        return render_template('admin_index.html')

    # 传形式化审核 审核信息 df  审核后新文档 docx_para
    filename = 'docxexam_userupload.docx'
    basepath = os.path.dirname(__file__)
    upload_path_svup2 = os.path.join(basepath, 'static/docx_repo/docx_gt/', filename)
    df = xsh_exam.compare(upload_path_svup2)

    filename = 'xsh_exam.docx'
    basepath = os.path.dirname(__file__)
    upload_path_xsh = os.path.join(basepath, 'static/docx_repo/docx_gt/', filename)
    docx_obj = Document(upload_path_xsh)
    docx_para = docx_obj.paragraphs
    dic = output_elements_from_docx(upload_path_svup2)
    # 合同要素金额大小写验证信息
    check_res_dict = check_res(dic)
    # 删除新生成的带标记的文档
    if os.path.exists(upload_path_xsh):
        os.remove(upload_path_xsh)
    # 删除生成的用户上传的文档
    if os.path.exists(upload_path_svup2):
        os.remove(upload_path_svup2)
    # 删除 kuploads 文件夹下的用户上传的文档
    if os.path.exists(upload_path_userup):
        os.remove(upload_path_userup)
    return render_template('admin_exam.html', docx_para=docx_para, df=df, dic=dic, check_res_dict=check_res_dict)


# 合同审核后的 要素表
@app.route('/admin_examdocx_ys', methods=['POST', 'GET'])
@login_required
def admin_examdocx_ys():
    if request.method == "POST":
        dic = []
        dic.append(request.values.get("docx_bh"))
        dic.append(request.values.get("docx_syy"))
        dic.append(request.values.get("docx_xyrq"))
        dic.append(request.values.get("docx_xydd"))
        dic.append(request.values.get("docx_zrf"))
        dic.append(request.values.get("docx_srf"))
        dic.append(request.values.get("docx_zwf"))
        dic.append(request.values.get("docx_zrffzr"))
        dic.append(request.values.get("docx_srffzr"))
        dic.append(request.values.get("docx_zwffzr"))
        dic.append(request.values.get("docx_zrfzs"))
        dic.append(request.values.get("docx_srfzs"))
        dic.append(request.values.get("docx_zwfzs"))
        dic.append(request.values.get("docx_zmbjye"))
        dic.append(request.values.get("docx_lx"))
        dic.append(request.values.get("docx_qtzq"))
        dic.append(request.values.get("docx_ztzq"))
        dic.append(request.values.get("docx_zrjk"))
        dic.append(request.values.get("docx_zqje"))
        dic.append(request.values.get("docx_wyj"))
        dic.append(request.values.get("docx_jzr"))
        dic.append(request.values.get("docx_bxze"))
        dic.append(request.values.get("docx_bjye"))
        dic.append(request.values.get("docx_qx"))
        dic.append(request.values.get("docx_khyh"))
        dic.append(request.values.get("docx_hm"))
        dic.append(request.values.get("docx_zh"))
        dic.append(request.values.get("docx_jybzz"))

        dic = d_x_dic(dic)
        docx_table.add_docx_table(dic)
        return render_template('admin_exam_ys.html', dic=dic)
    return render_template("admin_index.html")


# 合同审核后 选择合同流程表
@app.route('/admin_choosetable', methods=['POST', 'GET'])
@login_required
def admin_choosetable():
    if request.method == "POST":
        dic = []
        dic.append(request.values.get("doc_date"))
        dic.append(request.values.get("doc_loc"))
        return render_template("admin_choosetable.html", dic=dic)
    return render_template("admin_index.html")


# 合同审核后的合同流程表
@app.route('/admin_docx', methods=['POST', 'GET'])
@login_required
def admin_docx():
    if request.method == "POST":
        chooselist = request.values.getlist("cb")
        doc_date = request.values.get("doc_date")
        doc_loc = request.values.get("doc_loc")
        # 两个合同表
        if ("cd1" in chooselist and "cd2" in chooselist):
            return render_template("admin_docx.html", doc_date=doc_date, doc_loc=doc_loc)
        # 表1 合同面签记录表
        elif ("cd1" in chooselist and "cd2" not in chooselist):
            return render_template("admin_docx1.html", doc_date=doc_date, doc_loc=doc_loc)
        # 表2 合同审查申请表
        elif ("cd1" not in chooselist and "cd2" in chooselist):
            return render_template("admin_docx2.html")
        else:
            return render_template('admin_docx1.html', doc_date=doc_date, doc_loc=doc_loc)
    return render_template('admin_index.html')


# 保存保存保存保存合同面签记录表
@app.route('/admin_docx1_save', methods=['POST', 'GET'])
@login_required
def admin_docx1_save():
    if request.method == "POST":
        filename_e = 'docx_temp.docx'
        basepath_e = os.path.dirname(__file__)
        upload_path_temp_e = os.path.join(basepath_e, 'static/docx_repo/docx_gt/', filename_e)
        if os.path.exists(upload_path_temp_e):
            os.remove(upload_path_temp_e)
        pt_name = request.values.get("pt_name")
        name_1 = request.values.get("name_1")
        cont_1 = request.values.get("cont_1")
        othe_1 = request.values.get("othe_1")
        name_2 = request.values.get("name_2")
        cont_2 = request.values.get("cont_2")
        othe_2 = request.values.get("othe_2")
        name_3 = request.values.get("name_3")
        cont_3 = request.values.get("cont_3")
        othe_3 = request.values.get("othe_3")
        name_4 = request.values.get("name_4")
        cont_4 = request.values.get("cont_4")
        othe_4 = request.values.get("othe_4")
        name_5 = request.values.get("name_5")
        cont_5 = request.values.get("cont_5")
        othe_5 = request.values.get("othe_5")
        name_6 = request.values.get("name_6")
        cont_6 = request.values.get("cont_6")
        othe_6 = request.values.get("othe_6")
        name_7 = request.values.get("name_7")
        cont_7 = request.values.get("cont_7")
        othe_7 = request.values.get("othe_7")
        doc_date = request.values.get("doc_date")
        doc_loc = request.values.get("doc_loc")

        document = Document()
        document.add_heading("合同面签记录表", 0)
        document.add_paragraph('（我方人员填写）')
        document.add_heading('本记录在' + doc_date + "于" + doc_loc + "填写", level=2)
        document.add_paragraph('项目名称' + "       " + pt_name)
        table = document.add_table(rows=8, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '合同名称及编号'
        hdr_cells[1].text = '份数'
        hdr_cells[2].text = '备注'
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = name_1
        hdr_cells[1].text = cont_1
        hdr_cells[2].text = othe_1
        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = name_2
        hdr_cells[1].text = cont_2
        hdr_cells[2].text = othe_2
        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = name_3
        hdr_cells[1].text = cont_3
        hdr_cells[2].text = othe_3
        hdr_cells = table.rows[4].cells
        hdr_cells[0].text = name_4
        hdr_cells[1].text = cont_4
        hdr_cells[2].text = othe_4
        hdr_cells = table.rows[5].cells
        hdr_cells[0].text = name_5
        hdr_cells[1].text = cont_5
        hdr_cells[2].text = othe_5
        hdr_cells = table.rows[6].cells
        hdr_cells[0].text = name_6
        hdr_cells[1].text = cont_6
        hdr_cells[2].text = othe_6
        hdr_cells = table.rows[7].cells
        hdr_cells[0].text = name_7
        hdr_cells[1].text = cont_7
        hdr_cells[2].text = othe_7
        document.add_paragraph("面签见证人确认：上述合同由合同交易对手有权签字人当面签署，并以")
        document.add_paragraph("核对身份证原件及授权文件、企业营业执照原件及公章。")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("面签见证人签名：")
        document.add_paragraph("日期：")

        # 临时保存生成的文件
        filename = 'docx_temp.docx'
        basepath = os.path.dirname(__file__)
        upload_path_temp = os.path.join(basepath, 'static/docx_repo/docx_gt/', filename)
        document.save(upload_path_temp)

        upload_dic = os.path.join(basepath, 'static/docx_repo/docx_gt/')
        return send_from_directory(directory=upload_dic, filename="docx_temp.docx", as_attachment=True)

    return render_template("admin_index.html")


# 保存保存保存保存合同审查申请表
@app.route('/admin_docx2_save', methods=['POST', 'GET'])
@login_required
def admin_docx2_save():
    if request.method == "POST":
        filename_e = 'docx_temp2.docx'
        basepath_e = os.path.dirname(__file__)
        upload_path_temp_e = os.path.join(basepath_e, 'static/docx_repo/docx_gt/', filename_e)
        if os.path.exists(upload_path_temp_e):
            os.remove(upload_path_temp_e)
        dp_name = request.values.get("dp_name")
        pt_name = request.values.get("pt_name")
        pt_cont = request.values.get("pt_cont")
        xh = request.values.get("xh")
        mc = request.values.get("mc")
        tgf = request.values.get("tgf")
        wb = request.values.get("wb")
        xg = request.values.get("xg")
        yj = request.values.get("yj")
        sm = request.values.get("sm")
        qt = request.values.get("qt")
        document = Document()
        document.add_heading("合同审查申请表", 0)
        document.add_paragraph('申报部门：' + '    ' + dp_name)
        document.add_paragraph("")
        document.add_paragraph('项目名称：' + '    ' + pt_name)
        document.add_paragraph("")
        document.add_paragraph('项目编号：' + '    ' + pt_cont)
        document.add_paragraph("")
        table = document.add_table(rows=2, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '合同序号'
        hdr_cells[1].text = '合同名称'
        hdr_cells[2].text = '合同文本提供方'
        hdr_cells[3].text = '是否格式文本'
        hdr_cells[4].text = '是否允许修改'
        hdr_cells[5].text = '合同其他方意见'
        hdr_cells[6].text = '特别事项说明'
        hdr_cells[7].text = '其他'
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = xh
        hdr_cells[1].text = mc
        hdr_cells[2].text = tgf
        hdr_cells[3].text = wb
        hdr_cells[4].text = xg
        hdr_cells[5].text = yj
        hdr_cells[6].text = sm
        hdr_cells[7].text = qt
        document.add_paragraph("")
        document.add_paragraph("")

        # 临时保存生成的文件
        filename = 'docx_temp2.docx'
        basepath = os.path.dirname(__file__)
        upload_path_temp = os.path.join(basepath, 'static/docx_repo/docx_gt/', filename)
        document.save(upload_path_temp)
        upload_dic = os.path.join(basepath, 'static/docx_repo/docx_gt/')
        return send_from_directory(directory=upload_dic, filename="docx_temp2.docx", as_attachment=True)

    return render_template("admin_index.html")


# 合同审核后的合同流程表----审查申请表 还包含面签记录表
@app.route('/admin_docx4', methods=['POST', 'GET'])
@login_required
def admin_docx4():
    if request.method == "POST":
        doc_date = request.values.get("doc_date")
        doc_loc = request.values.get("doc_loc")
        return render_template('admin_docx4.html', doc_date=doc_date, doc_loc=doc_loc)
    return render_template('admin_index.html')


# 合同审核后的合同流程表----面签记录表
@app.route('/admin_docx1', methods=['POST', 'GET'])
@login_required
def admin_docx1():
    if request.method == "POST":
        doc_date = request.values.get("doc_date")
        doc_loc = request.values.get("doc_loc")
        return render_template('admin_docx1.html', doc_date=doc_date, doc_loc=doc_loc)
    return render_template('admin_index.html')


# 合同审核后的合同流程表----审查申请表
@app.route('/admin_docx2', methods=['POST', 'GET'])
@login_required
def admin_docx2():
    return render_template('admin_docx2.html')


# 合同生成  合同类型选择
@app.route('/admin_gena', methods=['POST', 'GET'])
@login_required
def admin_gena():
    return render_template('admin_gena.html')

elements_list = [
    '转让方',
    '转让方负责人',
    '转让方地址',
    '受让方',
    '受让方负责人',
    '受让方地址',
    '债务方',
    '债务方负责人',
    '债务方地址',
    '协议签订地点',
    '协议签订日期',
    '交易基准日',
    '债权本息总额',
    '本金余额',
    '欠息',
    '债权转让价款',
    '甲方指定账户',
    '开户银行',
    '户名',
    '账户',
    '交易保证金'
]
# 合同生成  合同信息填写
@app.route('/admin_gena2', methods=['POST', 'GET'])
@login_required
def admin_gena2():
    if request.method == "POST":
        template_class = request.values.get("xz_ejfl")
        # 债权转让协议_打包收购、债权转让协议_我司对外转让、债权转让协议_单户收购_双方签署
        choose_template_name = ''
        if template_class == '2_1':
            choose_template_name = '债权转让协议_我司对外转让'
        if template_class == '2_2':
            choose_template_name = '债权转让协议_打包收购'
        if template_class == '2_3':
            choose_template_name = '债权转让协议_单户收购_双方签署'
        if template_class == '2_4':
            choose_template_name = '债权转让协议_应收账款债权'

        csstip = [(x+1, 302 + x * 38) for x in range(len(get_elements_list()))]
        return render_template('admin_gena2.html',
                               csstip=csstip,
                               elements=elements_list,
                               choose_template_name=choose_template_name)
    return render_template('admin_gena.html')

# 合同生成  个性化条款生成
@app.route('/admin_gena3', methods=['POST', 'GET'])
@login_required
def admin_gena3():
    if request.method == 'POST':
        str1 = list()
        str1.append(request.values.get('ys_1'))
        str1.append(request.values.get('ys_2'))
        str1.append(request.values.get('ys_3'))
        str1.append(request.values.get('ys_4'))
        str1.append(request.values.get('ys_5'))
        str1.append(request.values.get('ys_6'))
        str1.append(request.values.get('ys_7'))
        str1.append(request.values.get('ys_8'))
        str1.append(request.values.get('ys_9'))
        str1.append(request.values.get('ys_10'))
        str1.append(request.values.get('ys_11'))
        str1.append(request.values.get('ys_12'))
        str1.append(strings_to_list(request.values.get('ys_13')))
        str1.append(strings_to_list(request.values.get('ys_14')))
        str1.append(strings_to_list(request.values.get('ys_15')))
        str1.append(strings_to_list(request.values.get('ys_16')))
        str1.append(request.values.get('ys_17'))
        str1.append(request.values.get('ys_18'))
        str1.append(request.values.get('ys_19'))
        str1.append(request.values.get('ys_20'))
        str1.append(strings_to_list(request.values.get('ys_21')))
        session['tmp_str1_list'] = str1  # 设置“字典”键值对
        choose_template_name = request.values.get("choose_template_name")
        return render_template('admin_gena3.html', choose_template_name=choose_template_name, str1=str1)
    return render_template('admin_gena.html')


# 合同生成  合同一键生成
@app.route('/admin_gena4', methods=['POST', 'GET'])
@login_required
def admin_gena4():
    if request.method == "POST":
        choose_template_name = request.values.get("choose_template_name")
        str1 = session.get('tmp_str1_list')
        # 将用户输入的信息填充到模板中 生成新的合同
        dic = list_to_dict(str1)
        contract_generation(choose_template_name, dic)
        project_path = os.path.dirname(__file__)
        template_path = os.path.join(project_path, 'static/docx_repo/docx_rd')
        fill_docx_path = os.path.join(template_path, 'saved_fill.docx')
        document = Document(fill_docx_path)
        content = document.paragraphs

        return render_template('admin_gena4.html', content=content)
    return render_template("admin_gena.html")


# 返回新生成的合同
@app.route('/admin_gena_newdocx', methods=['POST', 'GET'])
@login_required
def admin_gena_newdocx():
    if request.method == "POST":
        project_path = os.path.dirname(__file__)
        template_path = os.path.join(project_path, 'static/docx_repo/docx_gt')
        fill_docx_name = 'contract_generated.docx'
        return send_from_directory(directory=template_path, filename=fill_docx_name, as_attachment=True)
    return render_template("admin_gena.html")


# 意见反馈
@app.route('/admin_sugg',methods=['POST','GET'])
@login_required
def admin_sugg():
    if request.method == 'POST':
        suggest_text = request.form.get('demo')
        user_id = User.get_id(current_user)
        now_time = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        str_time = str(now_time)
        str_new = str_time[2:]
        sugg_id = int(str_new) + random.randint(0,5)
        sugg_id = sugg_id - 180000000000
        new_sugg = Sugg()
        new_sugg.sugg_id = sugg_id
        new_sugg.user_id = user_id
        new_sugg.sugg_text = suggest_text
        db.session.add(new_sugg)
        db.session.commit()
        suggest_all = db.session.query(Sugg).all()
        return render_template('admin_sugginfo.html',suggest_all = suggest_all)
    return render_template('admin_sugg.html')

# 意见反馈删除
@app.route('/admin_suggdel',methods=['POST','GET'])
@login_required
def admin_suggdel():
    if request.method == "POST":
        sugg_id = request.values.get("sugg_id")
        db.session.query(Sugg).filter_by(sugg_id=sugg_id).delete()
        db.session.commit()
        suggest_all = db.session.query(Sugg).all()
        db.session.commit()
        return render_template('admin_sugginfo.html', suggest_all=suggest_all)
    return render_template("admin_sugg.html")

# 系统管理
# 账户设置
@app.route('/admin_mage',methods=['POST','GET'])
@login_required
def admin_mage():
    return render_template('admin_mage.html')

# 系统信息
@app.route('/admin_info',methods=['POST','GET'])
@login_required
def admin_info():
    return render_template('admin_info.html')

# 登陆日志
@app.route('/admin_logo',methods=['POST','GET'])
@login_required
def admin_logo():
    return render_template('admin_logo.html')


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True)
    # app.run(debug=True)
