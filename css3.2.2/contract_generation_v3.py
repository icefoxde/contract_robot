# 函数名：contract_generation_v3
# 输入：用户选择的模板文档名称 用户填写的合同要素信息
# 输出：填充好的合同文档

import os
import re
from docx import Document
from docx.shared import *

# 四类合同模板文档所在路径
project_path = os.path.dirname(__file__)
template_path = os.path.join(project_path,'static/docx_repo/docx_rd')

template_docx_list = ['债权转让协议_单户收购_双方签署.docx', '债权转让协议_应收账款债权.docx', '债权转让协议_我司对外转让.docx', '债权转让协议_打包收购.docx']

txt_save_path = os.path.join(project_path,'static/docx_repo/docx_rd')

filled_docx_path = os.path.join(project_path,'static/docx_repo/docx_gt')


# 输入：一个模板文件名
# 输出：该模板文件中各合同要素所在段落 dict形式
def output_elements_para_number(template_filename):
    docx_text_string_list = []
    document = Document(template_filename)
    for i_para in document.paragraphs:
        i_para_text = i_para.text.strip()
        i_para_text = ''.join(i_para_text.split())
        docx_text_string_list.append(i_para_text)

    # 定义各类合同要素的段落号
    elements_para_number_dict = {
        '转让方':'',
        '转让方负责人': '',
        '转让方地址': '',
        '受让方': '',
        '受让方负责人': '',
        '受让方地址': '',
        '债务方': '',
        '债务方负责人': '',
        '债务方地址': '',
        '协议签订地点': '',
        '协议签订日期': '',
        '交易基准日': '',
        '债权本息总额': '',
        '本金余额': '',
        '欠息': '',
        '债权转让价款': '',
        '甲方指定账户': '',
        '开户银行': '',
        '户名': '',
        '账户': '',
        '交易保证金': ''
    }

    # 模板文档的段落 list 长度
    docx_string_list_lens = len(docx_text_string_list)

    # 提取转让方段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('转让方')
        j_res = string.find('甲方')
        if (z_res != -1 or j_res != -1) and elements_para_number_dict['转让方'] == '':
            elements_para_number_dict['转让方'] = str_i
            break


    # 提取转让方负责人段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        f_res = string.find('法定代表人')
        r_res = string.find('负责人')
        if (f_res != -1 or r_res != -1) and str_i-1>=0:
            string_pre = docx_text_string_list[str_i-1]
            z_res = string_pre.find('转让方')
            j_res = string_pre.find('甲方')
            if (z_res != -1 or j_res != -1) and elements_para_number_dict['转让方负责人'] == '':
                elements_para_number_dict['转让方负责人'] = str_i
                break



    # 提取转让方地址段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        dz_res = string.find('地址')
        zz_res = string.find('住址')
        zs_res = string.find('住所')
        if dz_res != -1 or zz_res != -1 or zs_res != -1 :
            if str_i-2 >= 0:
                string_pre = docx_text_string_list[str_i-2]
                j_res = string_pre.find('甲方')
                z_res = string_pre.find('转让方')
                if (j_res != -1 or z_res != -1) and elements_para_number_dict['转让方地址'] == '':
                    elements_para_number_dict['转让方地址'] = str_i
                    break



    # 提取受让方段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        s_res = string.find('受让方')
        y_res = string.find('乙方')
        if (s_res != -1 or y_res != -1) and elements_para_number_dict['受让方'] == '':
            elements_para_number_dict['受让方'] = str_i
            break


    # 提取受让方负责人段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        f_res = string.find('法定代表人')
        r_res = string.find('负责人')
        if (f_res != -1 or r_res != -1) and str_i-1>=0:
            string_pre = docx_text_string_list[str_i-1]
            s_res = string_pre.find('受让方')
            y_res = string_pre.find('乙方')
            if (s_res != -1 or y_res != -1) and elements_para_number_dict['受让方负责人'] == '':
                elements_para_number_dict['受让方负责人'] = str_i
                break


    # 提取受让方地址段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        dz_res = string.find('地址')
        zz_res = string.find('住址')
        zs_res = string.find('住所')
        if dz_res != -1 or zz_res != -1 or zs_res != -1 :
            if str_i-2 >= 0:
                string_pre = docx_text_string_list[str_i-2]
                y_res = string_pre.find('乙方')
                s_res = string_pre.find('受让方')
                if (y_res != -1 or s_res != -1) and elements_para_number_dict['受让方地址'] == '':
                    elements_para_number_dict['受让方地址'] = str_i
                    break


    # 提取债务方段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('债务方：')
        r_res = string.find('债务人：')
        if (z_res != -1 or r_res != -1) and elements_para_number_dict['债务方'] == '':
            elements_para_number_dict['债务方'] = str_i
            break


    # 提取债务方负责人段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        f_res = string.find('法定代表人')
        r_res = string.find('负责人')
        if (f_res != -1 or r_res != -1) and str_i-1>=0:
            string_pre = docx_text_string_list[str_i-1]
            z_res = string_pre.find('债务方：')
            r_res = string_pre.find('债务人：')
            if (z_res != -1 or r_res != -1) and elements_para_number_dict['债务方负责人'] == '':
                elements_para_number_dict['债务方负责人'] = str_i
                break


    # 提取债务方地址段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        dz_res = string.find('地址')
        zz_res = string.find('住址')
        zs_res = string.find('住所')
        if dz_res != -1 or zz_res != -1 or zs_res != -1 :
            if str_i-2 >= 0:
                string_pre = docx_text_string_list[str_i-2]
                r_res = string_pre.find('债务人：')
                z_res = string_pre.find('债务方：')
                if (r_res != -1 or z_res != -1) and elements_para_number_dict['债务方地址'] == '':
                    elements_para_number_dict['债务方地址'] = str_i
                    break


    # 提取协议签订地点段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        dz_res = string.find('协议签订地点：')
        if dz_res != -1 and elements_para_number_dict['协议签订地点'] == '':
            elements_para_number_dict['协议签订地点'] = str_i
            break


    # 提取协议签订日期段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        dz_res = string.find('协议签订日期：')
        if dz_res != -1 and elements_para_number_dict['协议签订日期'] == '':
            elements_para_number_dict['协议签订日期'] = str_i
            break



    # 提取交易基准日段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        j_res = string.find('基准日')
        n_res = string.find('年')
        y_res = string.find('月')
        r_res = string.find('日')
        if j_res != -1 and j_res < n_res and j_res < y_res and j_res < r_res and elements_para_number_dict['交易基准日'] == '':
            elements_para_number_dict['交易基准日'] = str_i
            break



    # 提取债权本息总额段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('债权本息总额')
        if z_res != -1 and elements_para_number_dict['债权本息总额'] == '':
            elements_para_number_dict['债权本息总额'] = str_i
            break


    # 提取本金余额段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('本金余额')
        if z_res != -1 and elements_para_number_dict['本金余额'] == '':
            elements_para_number_dict['本金余额'] = str_i
            break


    # 提取欠息段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('欠息')
        if z_res != -1 and elements_para_number_dict['欠息'] == '':
            elements_para_number_dict['欠息'] = str_i
            break


    # 提取债权转让价款段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('转让价款')
        r_res = string.find('人民币')
        y_res = string.find('元')
        if z_res != -1 and r_res != -1 and y_res != -1 and elements_para_number_dict['债权转让价款'] == '':
            elements_para_number_dict['债权转让价款'] = str_i
            break



    # 提取甲方指定账户
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('指定账户为')
        if z_res != -1 and elements_para_number_dict['甲方指定账户'] == '':
            elements_para_number_dict['甲方指定账户'] = str_i
            break


    # 提取开户银行段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('开户银行')
        if z_res != -1 and elements_para_number_dict['开户银行'] == '':
            elements_para_number_dict['开户银行'] = str_i
            break


    # 提取户名段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('户名')
        if z_res != -1 and elements_para_number_dict['户名'] == '':
            elements_para_number_dict['户名'] = str_i
            break


    # 提取账户段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('账户：')
        if z_res != -1 and elements_para_number_dict['账户'] == '':
            elements_para_number_dict['账户'] = str_i
            break


    # 提取交易保证金段落号
    for str_i in range(docx_string_list_lens):
        string = docx_text_string_list[str_i]
        z_res = string.find('交易保证金人民币')
        if z_res != -1 and elements_para_number_dict['交易保证金'] == '':
            elements_para_number_dict['交易保证金'] = str_i
            break


    for item in elements_para_number_dict:
        if elements_para_number_dict[item] == '':
            elements_para_number_dict[item] = -1

    return elements_para_number_dict


# 输入：四类模板文件路径
# 输出：每一类模板文档各合同要素段落号保存的txt文件
def save_para_num_txt():
    for item in template_docx_list:
        template_docx_path_i = os.path.join(template_path,item)
        elements_para_num_dict = output_elements_para_number(template_docx_path_i)
        txt_filename_i = item[:-5] + '.txt'

        # 将段落号写入txt文档中
        txt_file_path_i = os.path.join(txt_save_path,txt_filename_i)
        with open(txt_file_path_i,'w',encoding='utf-8') as f:
            for item in elements_para_num_dict:
                f.write(str(elements_para_num_dict[item]))
                f.write('\n')




# 输入：用户选择的模板文件名 用户输入的合同要素信息 dict形式
# 输出：填充的一个新的合同文档
def contract_generation(choose_template_name,elements_dict):

    # 拼接合同模板文档路径
    template_docx_path = ''
    # 拼接合同要素段落号txt路径
    template_txt_path = ''
    if choose_template_name == '债权转让协议_单户收购_双方签署':
        template_docx_filename = choose_template_name + '.docx'
        template_docx_path = os.path.join(template_path,template_docx_filename)
        template_txt_filename = choose_template_name + '.txt'
        template_txt_path = os.path.join(txt_save_path,template_txt_filename)

    if choose_template_name == '债权转让协议_应收账款债权':
        template_docx_filename = choose_template_name + '.docx'
        template_docx_path = os.path.join(template_path,template_docx_filename)
        template_txt_filename = choose_template_name + '.txt'
        template_txt_path = os.path.join(txt_save_path, template_txt_filename)

    if choose_template_name == '债权转让协议_我司对外转让':
        template_docx_filename = choose_template_name + '.docx'
        template_docx_path = os.path.join(template_path,template_docx_filename)
        template_txt_filename = choose_template_name + '.txt'
        template_txt_path = os.path.join(txt_save_path, template_txt_filename)

    if choose_template_name == '债权转让协议_打包收购':
        template_docx_filename = choose_template_name + '.docx'
        template_docx_path = os.path.join(template_path,template_docx_filename)
        template_txt_filename = choose_template_name + '.txt'
        template_txt_path = os.path.join(txt_save_path, template_txt_filename)

    # 开始对合同模板进行填充
    # 传入的模板文件路径存在
    if template_docx_path != '':
        document = Document(template_docx_path)

        # 将合同要素段落号转为dict
        elements_para_number_dict = {
            '转让方': '',
            '转让方负责人': '',
            '转让方地址': '',
            '受让方': '',
            '受让方负责人': '',
            '受让方地址': '',
            '债务方': '',
            '债务方负责人': '',
            '债务方地址': '',
            '协议签订地点': '',
            '协议签订日期': '',
            '交易基准日': '',
            '债权本息总额': '',
            '本金余额': '',
            '欠息': '',
            '债权转让价款': '',
            '甲方指定账户': '',
            '开户银行': '',
            '户名': '',
            '账户': '',
            '交易保证金': ''
        }
        with open(template_txt_path,'r',encoding='utf-8') as f:
            content = f.readlines()

        string = content[0]
        string = string[:-1]
        elements_para_number_dict['转让方'] = int(string)

        string = content[1]
        string = string[:-1]
        elements_para_number_dict['转让方负责人'] = int(string)

        string = content[2]
        string = string[:-1]
        elements_para_number_dict['转让方地址'] = int(string)

        string = content[3]
        string = string[:-1]
        elements_para_number_dict['受让方'] = int(string)

        string = content[4]
        string = string[:-1]
        elements_para_number_dict['受让方负责人'] = int(string)

        string = content[5]
        string = string[:-1]
        elements_para_number_dict['受让方地址'] = int(string)

        string = content[6]
        string = string[:-1]
        elements_para_number_dict['债务方'] = int(string)

        string = content[7]
        string = string[:-1]
        elements_para_number_dict['债务方负责人'] = int(string)

        string = content[8]
        string = string[:-1]
        elements_para_number_dict['债务方地址'] = int(string)

        string = content[9]
        string = string[:-1]
        elements_para_number_dict['协议签订地点'] = int(string)

        string = content[10]
        string = string[:-1]
        elements_para_number_dict['协议签订日期'] = int(string)

        string = content[11]
        string = string[:-1]
        elements_para_number_dict['交易基准日'] = int(string)

        string = content[12]
        string = string[:-1]
        elements_para_number_dict['债权本息总额'] = int(string)

        string = content[13]
        string = string[:-1]
        elements_para_number_dict['本金余额'] = int(string)

        string = content[14]
        string = string[:-1]
        elements_para_number_dict['欠息'] = int(string)

        string = content[15]
        string = string[:-1]
        elements_para_number_dict['债权转让价款'] = int(string)

        string = content[16]
        string = string[:-1]
        elements_para_number_dict['甲方指定账户'] = int(string)

        string = content[17]
        string = string[:-1]
        elements_para_number_dict['开户银行'] = int(string)

        string = content[18]
        string = string[:-1]
        elements_para_number_dict['户名'] = int(string)

        string = content[19]
        string = string[:-1]
        elements_para_number_dict['账户'] = int(string)

        string = content[20]
        string = string[:-1]
        elements_para_number_dict['交易保证金'] = int(string)


        # 填充转让方
        if elements_dict['转让方'] != '' and elements_para_number_dict['转让方'] != -1:
            # 用户输入的转让方信息
            fill_info = elements_dict['转让方']
            # 转让方所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['转让方']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接转让方string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填充转让方负责人
        if elements_dict['转让方负责人'] != '' and elements_para_number_dict['转让方负责人'] != -1:
            # 用户输入的转让方负责人信息
            fill_info = elements_dict['转让方负责人']
            # 转让方负责人所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['转让方负责人']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接转让方负责人string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写转让方地址
        if elements_dict['转让方地址'] != '' and elements_para_number_dict['转让方地址'] != -1:
            # 用户输入的转让方地址信息
            fill_info = elements_dict['转让方地址']
            # 转让方地址所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['转让方地址']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接转让方地址string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写受让方
        if elements_dict['受让方'] != '' and elements_para_number_dict['受让方'] != -1:
            # 用户输入的受让方信息
            fill_info = elements_dict['受让方']
            # 受让方所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['受让方']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接受让方string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写受让方负责人
        if elements_dict['受让方负责人'] != '' and elements_para_number_dict['受让方负责人'] != -1:
            # 用户输入的受让方负责人信息
            fill_info = elements_dict['受让方负责人']
            # 受让方负责人所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['受让方负责人']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接受让方负责人string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写受让方地址
        if elements_dict['受让方地址'] != '' and elements_para_number_dict['受让方地址'] != -1:
            # 用户输入的受让方地址信息
            fill_info = elements_dict['受让方地址']
            # 受让方地址所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['受让方地址']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接受让方地址string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname



        # 填写债务方
        if elements_dict['债务方'] != '' and elements_para_number_dict['债务方'] != -1:
            # 用户输入的债务方信息
            fill_info = elements_dict['债务方']
            # 债务方所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['债务方']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接债务方string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写债务方负责人
        if elements_dict['债务方负责人'] != '' and elements_para_number_dict['债务方负责人'] != -1:
            # 用户输入的债务方负责人信息
            fill_info = elements_dict['债务方负责人']
            # 债务方负责人所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['债务方负责人']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接债务方负责人string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写债务方地址
        if elements_dict['债务方地址'] != '' and elements_para_number_dict['债务方地址'] != -1:
            # 用户输入的债务方地址信息
            fill_info = elements_dict['债务方地址']
            # 债务方地址所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['债务方地址']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接债务方地址string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写协议签订地点
        if elements_dict['协议签订地点'] != '' and elements_para_number_dict['协议签订地点'] != -1:
            # 用户输入的协议签订地点信息
            fill_info = elements_dict['协议签订地点']
            # 协议签订地点所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['协议签订地点']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接协议签订地点string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写协议签订日期
        if elements_dict['协议签订日期'] != '' and elements_para_number_dict['协议签订日期'] != -1:
            # 用户输入的协议签订日期信息
            fill_info = elements_dict['协议签订日期']
            # 协议签订日期所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['协议签订日期']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接协议签订日期string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写甲方指定账户
        if elements_dict['甲方指定账户'] != '' and elements_para_number_dict['甲方指定账户'] != -1:
            # 用户输入的甲方指定账户信息
            fill_info = elements_dict['甲方指定账户']
            # 甲方指定账户所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['甲方指定账户']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接甲方指定账户string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写开户银行
        if elements_dict['开户银行'] != '' and elements_para_number_dict['开户银行'] != -1:
            # 用户输入的开户银行信息
            fill_info = elements_dict['开户银行']
            # 开户银行所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['开户银行']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接开户银行string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写户名
        if elements_dict['户名'] != '' and elements_para_number_dict['户名'] != -1:
            # 用户输入的户名信息
            fill_info = elements_dict['户名']
            # 户名所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['户名']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接户名string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写账户
        if elements_dict['账户'] != '' and elements_para_number_dict['账户'] != -1:
            # 用户输入的账户信息
            fill_info = elements_dict['账户']
            # 账户所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['账户']]
            template_string = para_element.text
            loc = template_string.find('：')
            # 拼接账户string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写交易基准日
        if elements_dict['交易基准日'] != '' and elements_para_number_dict['交易基准日'] != -1:
            # 用户输入的交易基准日信息
            fill_info = elements_dict['交易基准日']
            # 交易基准日所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['交易基准日']]
            template_string = para_element.text
            loc = template_string.find('即')
            # 拼接交易基准日string
            template_string_filled = template_string[:loc+1] + fill_info

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写债权本息总额、本金余额、欠息
        if elements_para_number_dict['债权本息总额'] != -1:

            zq_dx_string = elements_dict['债权本息总额'][0]
            zq_xx_string = elements_dict['债权本息总额'][1]

            bj_dx_string = elements_dict['本金余额'][0]
            bj_xx_string = elements_dict['本金余额'][1]

            qx_dx_string = elements_dict['欠息'][0]
            qx_xx_string = elements_dict['欠息'][1]

            para_element = document.paragraphs[elements_para_number_dict['债权本息总额']]
            template_string = para_element.text
            loc = template_string.find('币')
            template_string_filled = template_string[:loc + 1] + zq_dx_string + '元（小写：人民币' + zq_xx_string + '元），' + '其中本金余额人民币' + bj_dx_string + '元（小写：人民币' + bj_xx_string + '元），' + '欠息人民币' + qx_dx_string + '元（小写：人民币' + qx_xx_string + '元）。'

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 填写债权转让价款
        if elements_dict['债权转让价款'] != '' and elements_para_number_dict['债权转让价款'] != -1:
            # 用户输入的债权转让价款信息
            zqzr_dx_string = elements_dict['债权转让价款'][0]
            zqzr_xx_string = elements_dict['债权转让价款'][1]

            # 债权转让价款所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['债权转让价款']]
            template_string = para_element.text
            loc = template_string.find('币')
            # 拼接债权转让价款string
            template_string_filled = template_string[:loc+1] + zqzr_dx_string + '元（小写：人民币' + zqzr_xx_string + '元）。'

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname



        # 填写交易保证金
        if elements_dict['交易保证金'] != '' and elements_para_number_dict['交易保证金'] != -1:
            # 用户输入的交易保证金信息
            jy_dx_string = elements_dict['交易保证金'][0]
            jy_xx_string = elements_dict['交易保证金'][1]

            # 交易保证金所在段落文本
            para_element = document.paragraphs[elements_para_number_dict['交易保证金']]
            template_string = para_element.text
            loc = template_string.find('币')
            # 拼接交易保证金string
            template_string_filled = template_string[:loc+1] + jy_dx_string + '元（小写：人民币' + jy_xx_string + '元）。'

            fontsize = 14.0
            fontname = u'宋体'
            for i in range(len(para_element.runs)):
                para_element.runs[i].clear()
            run1 = para_element.add_run(template_string_filled)
            run1.font.size = Pt(fontsize)
            run1.font.name = fontname


        # 保存填充后的合同文档
        filled_docx_filename = 'contract_generated.docx'
        save_filled_docx_path = os.path.join(filled_docx_path,filled_docx_filename)
        document.save(save_filled_docx_path)


if __name__ == '__main__':
    dic={
        '转让方': '',
        '转让方负责人': '',
        '转让方地址': '',
        '受让方': '',
        '受让方负责人': '',
        '受让方地址': '',
        '债务方': '',
        '债务方负责人': '',
        '债务方地址': '',
        '协议签订地点': '',
        '协议签订日期': '',
        '交易基准日': '',
        '债权本息总额': ['壹佰', '100'],
        '本金余额': ['壹佰', '100'],
        '欠息': ['壹佰', '100'],
        '债权转让价款': ['壹佰', '100'],
        '甲方指定账户': '',
        '开户银行': '',
        '户名': '',
        '账户': '',
        '交易保证金': ['壹佰', '100']
    }
    # filename_list = ['债权转让协议_单户收购_双方签署', '债权转让协议_应收账款债权', '债权转让协议_我司对外转让', '债权转让协议_打包收购']
    # for item in filename_list:
    #     contract_generation(item, dic)
    save_para_num_txt()



